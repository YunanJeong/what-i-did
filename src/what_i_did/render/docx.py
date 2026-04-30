from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from what_i_did import prompts
from what_i_did.models import PortfolioSummary, RepoAnalysis


def _add_hyperlink(paragraph, url: str, text: str):
    # python-docx 는 하이퍼링크 API 를 공개하지 않으므로 OOXML 요소를 직접 구성한다.
    # 문서의 relationships 에 외부 URL 을 등록(part.relate_to)하고, 그 r_id 를 참조하는
    # <w:hyperlink> 를 paragraph 에 삽입하는 패턴.
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_repo_card(doc, a: RepoAnalysis, p) -> None:
    h = doc.add_heading(level=3)
    _add_hyperlink(h, a.html_url, a.name)

    meta_bits = []
    if a.primary_language:
        meta_bits.append(a.primary_language)
    if a.stars:
        meta_bits.append(f"{p.LABEL_STARS}: {a.stars}")
    if meta_bits:
        meta_par = doc.add_paragraph()
        meta_run = meta_par.add_run(" · ".join(meta_bits))
        meta_run.italic = True

    if a.description:
        quote = doc.add_paragraph(a.description)
        quote.paragraph_format.left_indent = quote.paragraph_format.left_indent

    if a.purpose:
        par = doc.add_paragraph()
        par.add_run(f"{p.LABEL_PURPOSE}: ").bold = True
        par.add_run(a.purpose)

    if a.key_features:
        par = doc.add_paragraph()
        par.add_run(f"{p.LABEL_FEATURES}:").bold = True
        _add_bullets(doc, a.key_features)

    if a.tech_stack:
        par = doc.add_paragraph()
        par.add_run(f"{p.LABEL_TECH}: ").bold = True
        par.add_run(", ".join(a.tech_stack))

    if a.highlights:
        par = doc.add_paragraph()
        par.add_run(f"{p.LABEL_HIGHLIGHTS}:").bold = True
        _add_bullets(doc, a.highlights)

    if a.user_contribution_notes:
        par = doc.add_paragraph()
        par.add_run(f"{p.LABEL_CONTRIB}: ").bold = True
        par.add_run(a.user_contribution_notes)


def render(portfolio: PortfolioSummary, out_path: Path) -> None:
    p = prompts.get(portfolio.lang)
    doc = Document()
    doc.add_heading(p.DOC_TITLE.format(username=portfolio.username), level=0)

    if portfolio.overall_summary:
        doc.add_heading(p.SEC_SUMMARY, level=1)
        doc.add_paragraph(portfolio.overall_summary)

    if portfolio.main_projects:
        doc.add_heading(p.SEC_MAIN, level=1)
        for a in portfolio.main_projects:
            _add_repo_card(doc, a, p)

    if portfolio.side_projects:
        doc.add_heading(p.SEC_SIDE, level=1)
        for a in portfolio.side_projects:
            _add_repo_card(doc, a, p)

    if portfolio.tech_counts:
        doc.add_heading(p.SEC_STACK, level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid"
        hdr = table.rows[0].cells
        hdr[0].text = p.LABEL_TECH
        hdr[1].text = "#"
        for tech, count in portfolio.tech_counts:
            row = table.add_row().cells
            row[0].text = tech
            row[1].text = str(count)

    footer = doc.add_paragraph()
    run = footer.add_run(
        f"{p.LABEL_GENERATED}: {portfolio.generated_at.strftime('%Y-%m-%d %H:%M UTC')} · "
        f"{p.LABEL_MODEL}: {portfolio.model}"
    )
    run.italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
